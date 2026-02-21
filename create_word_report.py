#!/usr/bin/env python3
"""
生成包含图片和文字报告的 Word 文档
"""

import os
import glob
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_report(image_path: str, text_report: str, output_path: str):
    """创建包含图片和文字的 Word 报告"""

    # 创建文档
    doc = Document()

    # 添加标题
    title = doc.add_heading('BTC 交易分析报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加报告时间
    doc.add_paragraph(f'报告时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')

    # 添加第一部分标题
    doc.add_heading('📊 市场分析', 1)

    # 添加图片
    if os.path.exists(image_path):
        # 添加图片说明
        doc.add_paragraph('📈 净值曲线图')

        # 插入图片
        try:
            doc.add_picture(image_path, width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f'图片插入失败: {str(e)}')
    else:
        doc.add_paragraph('图片文件不存在')

    doc.add_paragraph('')

    # 添加第二部分标题
    doc.add_heading('📋 交易回测报告', 1)

    # 添加文字报告
    # 分段添加文字
    lines = text_report.split('\n')
    for line in lines:
        if line.startswith('══') or line.startswith('═'):
            # 添加分隔线
            p = doc.add_paragraph(line)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('###') or line.startswith('#'):
            # 添加标题
            heading_text = line.lstrip('#').strip()
            level = 1 if line.startswith('###') else 2
            doc.add_heading(heading_text, level)
        elif line.strip() == '':
            # 空行
            doc.add_paragraph()
        else:
            # 普通段落
            doc.add_paragraph(line)

    # 保存文档
    try:
        doc.save(output_path)
        return True, f"Word 文档已保存: {output_path}"
    except Exception as e:
        return False, f"保存失败: {str(e)}"

def generate_latest_word_report():
    """生成最新的 Word 报告"""
    # 路径
    image_path = "/root/.openclaw/workspace/backtest_chart.png"
    ascii_chart_path = "/root/.openclaw/workspace/ascii_line_chart.txt"
    output_dir = "/root/.openclaw/workspace"

    # 读取文字报告
    text_report = ""

    # 优先读取 ASCII 图表
    if os.path.exists(ascii_chart_path):
        with open(ascii_chart_path, 'r', encoding='utf-8') as f:
            text_report = f.read()

    # 读取市场分析报告
    market_report_files = sorted(glob.glob('/root/.openclaw/workspace/reports/btc_report_*.txt'))
    if market_report_files:
        with open(market_report_files[-1], 'r', encoding='utf-8') as f:
            market_report = f.read()
        text_report = f"{market_report}\n\n{text_report}"

    # 读取回测报告
    backtest_report_files = sorted(glob.glob('/root/.openclaw/workspace/reports/btc_backtest_report_*.txt'))
    if backtest_report_files:
        with open(backtest_report_files[-1], 'r', encoding='utf-8') as f:
            backtest_report = f.read()
        text_report = f"{text_report}\n\n{backtest_report}"

    # 生成输出文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(output_dir, f"btc_report_{timestamp}.docx")

    # 创建 Word 文档
    success, message = create_word_report(image_path, text_report, output_path)

    return success, message, output_path

if __name__ == "__main__":
    success, message, output_path = generate_latest_word_report()
    print(message)
    if success:
        print(f"文档路径: {output_path}")
